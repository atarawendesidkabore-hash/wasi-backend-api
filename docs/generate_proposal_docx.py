"""
Generate a professional Word document for the WASI Government Proposal.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style definitions ───────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B3A5C"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            # Alternate row shading
            if r_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF0F7"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table


def add_bold_paragraph(doc, text, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


def add_paragraph(doc, text, bold_prefix=None, size=11):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


# ═══════════════════════════════════════════════════════════════
#                        COVER PAGE
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

add_bold_paragraph(doc, "BUSINESS PROPOSAL", size=28,
                   color=RGBColor(0x1B, 0x3A, 0x5C),
                   align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_bold_paragraph(doc, "National Sovereign Data Tokenization Program\nfor Burkina Faso",
                   size=18, color=RGBColor(0x1B, 0x3A, 0x5C),
                   align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_horizontal_line(doc)
doc.add_paragraph()

cover_lines = [
    ("Submitted to:", True),
    ("His Excellency Captain Ibrahim TRAORÉ, President of Faso, Head of State", False),
    ("The Prime Minister, Head of Government", False),
    ("The Minister of Digital Transition, Posts and Telecommunications", False),
    ("The Minister of Economy and Finance", False),
    ("", False),
    ("Submitted by: WASI — West African Shipping & Economic Intelligence", True),
    ("Date: March 2026", True),
    ("Reference: WASI/BF/GOV/2026-001", True),
]

for text, bold in cover_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#                       COVER LETTER
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Cover Letter', level=1)
add_horizontal_line(doc)

add_paragraph(doc, "Your Excellency, Honorable Ministers,")
doc.add_paragraph()

cover_paragraphs = [
    "I have the honor of submitting to your esteemed attention the present proposal concerning the creation of a National Sovereign Data Tokenization Program for Burkina Faso.",

    "This program aims to transform every economic activity of every Burkinabè citizen — from the millet farmer in Djibo to the cement trader in Bobo-Dioulasso — into structured, verifiable, and commercially viable data that can be sold on global artificial intelligence and economic analysis markets.",

    "The program rests on three pillars:",
]

for text in cover_paragraphs:
    add_paragraph(doc, text)

pillars = [
    ("1. Citizen Data Income", " — enabling every Burkinabè to earn a daily income by simply declaring their activities via USSD (*384*WASI#), requiring no smartphone, no internet, and no literacy."),
    ("2. Data Tokenization Tax Credit (DTTC)", " — incentivizing formal and informal businesses to share their operational data in exchange for tax reductions, thereby formalizing the invisible economy."),
    ("3. Digital Faso Meabo", " — accelerating the execution of public contracts through tokenized work milestones, real-time citizen verification, and direct payment to community laborers."),
]

for bold_part, rest in pillars:
    add_bullet(doc, rest, bold_prefix=bold_part)

doc.add_paragraph()

closing_paragraphs = [
    "The stakes go beyond technology. This is about economic sovereignty, massive job creation in fragile zones, and fighting terrorism through economic inclusion. A young man earning 15,000 CFA per month by declaring his data will not join an armed group offering the same amount at the cost of his life.",

    "Burkina Faso would become the first African country to monetize its sovereign data for the direct benefit of its population. A historic achievement that His Excellency the President of Faso could announce on the international stage.",

    "We request the opportunity to present this program in detail before your technical services.",

    "Please accept, Your Excellency, Honorable Ministers, the expression of my highest consideration.",
]

for text in closing_paragraphs:
    add_paragraph(doc, text)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("[Name of CEO]")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Calibri'
add_paragraph(doc, "WASI — West African Shipping & Economic Intelligence")
add_paragraph(doc, "Ouagadougou, Burkina Faso")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#                     TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
add_horizontal_line(doc)

toc_items = [
    "1. Executive Summary",
    "2. The Problem: Burkina Faso's Invisible Economy",
    "3. The Solution: Sovereign Data Tokenization",
    "4. Pillar 1: Citizen Data Income",
    "5. Pillar 2: Data Tokenization Tax Credit (DTTC)",
    "6. Pillar 3: Digital Faso Meabo",
    "7. National Security Impact",
    "8. Technical Architecture",
    "9. Economic Model and Financial Projections",
    "10. Proposed Legal Framework",
    "11. Implementation Timeline",
    "12. International Partnerships",
    "13. Annexes",
]

for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#                  1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════

doc.add_heading('1. Executive Summary', level=1)
add_horizontal_line(doc)

doc.add_heading('The Problem', level=2)

add_paragraph(doc, "Burkina Faso possesses considerable economic wealth that remains invisible. Approximately 35 to 40% of GDP belongs to the informal economy: millions of daily transactions, harvests, cross-border trade flows, and fund movements that escape all measurement. This invisibility deprives the State of tax revenue, deprives citizens of access to credit, and deprives policymakers of reliable data for effective public policy.")

add_paragraph(doc, "Simultaneously, the global artificial intelligence economy has created unprecedented demand for structured African data. AI companies (Anthropic, OpenAI, Google), financial institutions, humanitarian organizations, and commodity traders are willing to pay millions of dollars for reliable, fresh, and verifiable African data. This data currently exists nowhere in exploitable form.")

doc.add_heading('The Solution', level=2)

add_paragraph(doc, "We propose a National Sovereign Data Tokenization Program that:")

solution_items = [
    "Transforms every citizen's economic activity into a tradeable data token",
    "Directly compensates citizens via mobile money for each declared data point",
    "Grants tax credits to businesses that share their operational data",
    "Accelerates public contracts through real-time monitoring and direct community payment",
    "Sells aggregated, anonymized data on global markets",
    "Generates sovereign revenue for the State while reducing poverty and terrorism",
]

for item in solution_items:
    add_bullet(doc, item)

doc.add_heading('Key Figures', level=2)

add_table(doc,
    ["Indicator", "Value"],
    [
        ["Annual revenue for the State (Year 3)", "4,800,000,000 CFA (~$7.4M)"],
        ["Direct jobs created (Village Data Chiefs)", "8,000"],
        ["Citizens compensated for their data", "200,000"],
        ["Businesses formalized through DTTC", "15,000"],
        ["Estimated reduction in terrorist recruitment", "10–25% in covered zones"],
        ["Initial investment required", "3,200,000,000 CFA (~$4.9M)"],
        ["Return on investment", "18 months"],
    ],
    col_widths=[8, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#       2. THE PROBLEM: BURKINA FASO'S INVISIBLE ECONOMY
# ═══════════════════════════════════════════════════════════════

doc.add_heading("2. The Problem: Burkina Faso's Invisible Economy", level=1)
add_horizontal_line(doc)

doc.add_heading('2.1 The Informal Economy', level=2)

add_paragraph(doc, "Burkina Faso has approximately 22 million inhabitants. Among them:")

informal_facts = [
    ("80%", " live from agriculture, herding, or small trade"),
    ("90%", " of businesses are informal (without a tax identification number — NIF/IFU)"),
    ("65%", " of cross-border commercial exchanges are undeclared"),
    ("70%", " of transactions are conducted in cash, leaving no digital trace"),
]
for bold, rest in informal_facts:
    add_bullet(doc, rest, bold_prefix=bold)

add_paragraph(doc, "These activities represent real economic value — billions of CFA per day — but they are invisible to the State, to banks, and to international partners.")

doc.add_heading('2.2 The Consequences of Invisibility', level=2)

add_bold_paragraph(doc, "For the State:", size=12)
add_bullet(doc, "Reduced tax base: the State cannot tax what it cannot see")
add_bullet(doc, "Blind planning: agricultural policies rely on data that is 2 to 5 years old")
add_bullet(doc, "Aid dependency: without reliable data, technical and financial partners impose their own surveys")

add_bold_paragraph(doc, "For Citizens:", size=12)
add_bullet(doc, "No financial history → no access to bank credit")
add_bullet(doc, "No proof of activity → no social protection")
add_bullet(doc, "No visibility → feeling of abandonment by the State → vulnerability to radicalization")

add_bold_paragraph(doc, "For Public Contracts:", size=12)
add_bullet(doc, "No real-time monitoring of contract execution")
add_bullet(doc, "Contractors delay work without consequence")
add_bullet(doc, "Average cost overrun on public contracts is estimated at 30% of the initial amount")
add_bullet(doc, "Beneficiary communities have no voice in monitoring projects that concern them")

doc.add_heading('2.3 The Global Opportunity', level=2)

add_paragraph(doc, "The AI economy has created a new market:")

add_table(doc,
    ["Potential Buyer", "What They Seek", "Estimated Annual Budget"],
    [
        ["AI Companies (Anthropic, OpenAI, Google)", "Training data on Africa", "$50,000–$500,000/year"],
        ["Fintech (Wave, Orange Money, Ecobank)", "Credit risk data", "$10,000–$100,000/year"],
        ["Commodity Traders (Cargill, Olam)", "Real-time pricing & logistics", "$20,000–$200,000/year"],
        ["International Orgs (UN, World Bank)", "Real-time economic indicators", "$5,000–$50,000/year/project"],
        ["ECOWAS Governments", "Cross-border trade data", "Revenue sharing"],
    ],
    col_widths=[5.5, 5, 5.5]
)

add_paragraph(doc, "These buyers currently pay for mediocre, years-old data. Fresh, verified data produced by economic actors themselves would command a premium price.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#        3. THE SOLUTION: SOVEREIGN DATA TOKENIZATION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('3. The Solution: Sovereign Data Tokenization', level=1)
add_horizontal_line(doc)

doc.add_heading('3.1 What is Data Tokenization?', level=2)

add_paragraph(doc, 'Data tokenization consists of transforming each economic activity into a structured data unit (a "token") that can be:')

token_steps = [
    ("Collected", " — via USSD on a basic phone (no smartphone needed)"),
    ("Verified", " — by cross-referencing with other sources (citizens, businesses, sensors)"),
    ("Aggregated", " — into anonymized indicators (no personal data is sold)"),
    ("Commercialized", " — sold to global buyers via a secure API"),
    ("Compensated", " — the data producer receives an immediate payment via mobile money"),
]
for i, (bold, rest) in enumerate(token_steps):
    add_bullet(doc, rest, bold_prefix=f"{i+1}. {bold}")

doc.add_heading('3.2 The Fundamental Principle', level=2)

principle_text = (
    "The farmer who reports the millet price at the Kaya market → produces a data token → "
    "receives 50 CFA on Orange Money → that data point is sold for 65 CFA to an international buyer → "
    "the State receives 20 CFA as a sovereign royalty → Burkina Faso grows wealthier from the intelligence of its own people."
)
p = add_paragraph(doc, principle_text)
# Add a colored background to make it stand out
for run in p.runs:
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_heading('3.3 The Three Pillars', level=2)

add_table(doc,
    ["Pillar", "Mechanism", "Target", "Impact"],
    [
        ["1. Citizen Data Income", "Citizens declare daily activities via USSD and receive instant mobile money payment", "200,000 citizens", "Poverty reduction + security"],
        ["2. Data Tokenization Tax Credit (DTTC)", "Businesses share operational data for tax reductions", "15,000 businesses", "Formalization of the informal economy"],
        ["3. Digital Faso Meabo", "Public contracts monitored in real time; community labor directly compensated", "All contracts > 50M CFA", "30% less delays, 15–25% cost savings"],
    ],
    col_widths=[3.5, 5, 3, 4.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#            4. PILLAR 1: CITIZEN DATA INCOME
# ═══════════════════════════════════════════════════════════════

doc.add_heading('4. Pillar 1: Citizen Data Income', level=1)
add_horizontal_line(doc)

doc.add_heading('4.1 Concept', level=2)

add_paragraph(doc, "Every Burkinabè citizen carries valuable economic information in their head: the price of millet at the local market, the state of the road to the nearest town, their harvest volume, fuel availability. This information has global market value, yet citizens are never compensated for it.")

add_paragraph(doc, "The Citizen Data Income allows any Burkinabè to earn a daily income by simply declaring what they know, what they do, and what they observe.")

doc.add_heading('4.2 How It Works', level=2)

add_bold_paragraph(doc, "Registration", size=12)
add_paragraph(doc, "The citizen dials *384*WASI# on any phone (even a basic 2,000 CFA handset). They provide their Orange Money/Wave number and location. Registration takes 60 seconds. No identity document is required to begin. Phone numbers are hashed in SHA-256 to protect privacy.")

add_bold_paragraph(doc, "Daily Declaration Menu", size=12)

add_table(doc,
    ["Option", "Declaration Type", "Payment"],
    [
        ["1", "Report a market price (millet, maize, rice, livestock...)", "50 CFA"],
        ["2", "Declare a harvest or agricultural work", "100 CFA"],
        ["3", "Report road conditions", "75 CFA"],
        ["4", "Declare a border crossing (with goods)", "150 CFA"],
        ["5", "Report fuel price", "50 CFA"],
        ["6", "Report rainfall", "50 CFA"],
        ["7", "Declare today's activity (even \"nothing\")", "200 CFA"],
    ],
    col_widths=[2, 10, 3]
)

add_bold_paragraph(doc, "Instant Payment", size=12)
add_paragraph(doc, "For each validated declaration, the corresponding amount is credited immediately to the citizen's mobile money account. No waiting, no paperwork, no middlemen.")

doc.add_heading('4.3 Potential Earnings', level=2)

add_bold_paragraph(doc, "Scenario: Farmer in Soum Province", size=12)

add_table(doc,
    ["Time", "Activity", "Earned"],
    [
        ["Morning", "Declares today's activity (agricultural work)", "200 CFA"],
        ["Morning", "Reports millet price at Djibo market", "50 CFA"],
        ["Midday", "Reports road condition to Ouagadougou", "75 CFA"],
        ["Evening", "Reports diesel price at local station", "50 CFA"],
        ["", "Daily total", "375 CFA"],
        ["", "Monthly total (26 working days)", "9,750 CFA (~$15)"],
    ],
    col_widths=[3, 8.5, 3.5]
)

add_paragraph(doc, "9,750 CFA/month represents 28% of the minimum wage — earned in 5 minutes of daily work via a basic phone. This covers school fees for 2 children or one month's grain for a household.")

doc.add_heading('4.4 Career Progression', level=2)

add_table(doc,
    ["Level", "Monthly Income", "Requirements"],
    [
        ["Basic reporter", "6,000–10,000 CFA", "Regular daily declarations"],
        ["Active reporter", "10,000–16,000 CFA", "Multiple declaration types, good reliability"],
        ["Village Data Chief", "40,000–60,000 CFA", "Trains and supervises 20+ reporters"],
        ["Provincial Coordinator", "80,000–120,000 CFA", "Manages 10+ villages"],
    ],
    col_widths=[4.5, 4, 7.5]
)

doc.add_heading('4.5 The Daily Activity Declaration as "Proof of Life"', level=2)

add_paragraph(doc, 'The daily activity declaration (option 7) is the most valuable data point in the system, even when the citizen declares doing nothing.')

add_bold_paragraph(doc, "Stability Indicator:", size=11)
add_paragraph(doc, "If 2,400 registered participants in Soum Province submit 2,350 reports (97.9%) → NORMAL, zone is stable.")

add_bold_paragraph(doc, "Early Warning:", size=11)
add_paragraph(doc, "If Seno Province drops from 1,790 reports (99.4%) yesterday to 340 reports (18.9%) today → ALERT. 81% drop signals possible displacement or security incident. Alert sent automatically to CONASUR, UN OCHA, and Defense Forces.")

p = add_paragraph(doc, "The absence of data IS data. A sudden drop in declarations from a zone is the fastest early warning signal that exists — faster than satellites, faster than military intelligence, faster than the press.")
for run in p.runs:
    run.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#      5. PILLAR 2: DATA TOKENIZATION TAX CREDIT (DTTC)
# ═══════════════════════════════════════════════════════════════

doc.add_heading('5. Pillar 2: Data Tokenization Tax Credit (DTTC)', level=1)
add_horizontal_line(doc)

doc.add_heading('5.1 Concept', level=2)

add_paragraph(doc, "Burkinabè businesses hold valuable operational data (sales, inventory, purchases, employees, suppliers). Currently, they have no incentive to share this data. The DTTC reverses this logic: a business that shares structured data receives a tax credit proportional to the value of the data provided.")

doc.add_heading('5.2 Tax Mechanism', level=2)

add_table(doc,
    ["Item", "Amount"],
    [
        ["Tax owed by business", "2,000,000 CFA"],
        ["Tokenized data contributed (calculated value)", "500,000 CFA"],
        ["Tax paid in cash", "1,500,000 CFA"],
        ["", ""],
        ["The State 'loses' in tax revenue", "500,000 CFA"],
        ["The State SELLS that data for", "1,200,000 CFA"],
        ["NET GAIN FOR THE STATE", "+700,000 CFA"],
    ],
    col_widths=[10, 5]
)

doc.add_heading('5.3 Credit Schedule by Data Type', level=2)

add_bold_paragraph(doc, "Tier A — Operational Data (Cap: 1,300,000 CFA/year)", size=12)

add_table(doc,
    ["Data", "Frequency", "Annual Token Value"],
    [
        ["Revenue by product line", "Monthly", "300,000 CFA"],
        ["Headcount and payroll", "Quarterly", "200,000 CFA"],
        ["Supplier payments and origins", "Monthly", "250,000 CFA"],
        ["Inventory levels", "Weekly", "400,000 CFA"],
        ["Energy consumption (kWh)", "Monthly", "150,000 CFA"],
    ],
    col_widths=[6, 4, 5]
)

add_bold_paragraph(doc, "Tier B — Market Intelligence (Cap: 750,000 CFA/year)", size=12)

add_table(doc,
    ["Data", "Frequency", "Annual Token Value"],
    [
        ["Customer transaction volumes", "Monthly", "200,000 CFA"],
        ["Input prices (raw materials)", "Weekly", "300,000 CFA"],
        ["Transport/logistics costs", "Monthly", "150,000 CFA"],
        ["Equipment/machinery status", "Quarterly", "100,000 CFA"],
    ],
    col_widths=[6, 4, 5]
)

add_bold_paragraph(doc, "Tier C — Strategic Data (Cap: 1,800,000 CFA/year)", size=12)

add_table(doc,
    ["Data", "Frequency", "Annual Token Value"],
    [
        ["Cross-border trade invoices", "Per transaction", "500,000 CFA"],
        ["Mining output volumes", "Monthly", "600,000 CFA"],
        ["Agricultural yield per hectare", "Seasonal", "400,000 CFA"],
        ["Construction project milestones", "Quarterly", "300,000 CFA"],
    ],
    col_widths=[6, 4, 5]
)

add_paragraph(doc, "Overall cap: The DTTC cannot exceed 25% of the tax owed, with an absolute cap of 5,000,000 CFA per business per year. The State always collects at least 75% in cash.")

doc.add_heading('5.4 The Formalization Flywheel', level=2)

flywheel_steps = [
    "Business wants the tax credit",
    "To qualify, business must have a tax ID (NIF/IFU) → business is now FORMAL",
    "Formal business gains access to bank credit",
    "Bank uses WASI data to score the business → loan granted → business grows and hires",
    "More employees = more tax revenue + more citizen data reporters",
    "Tax revenue INCREASES despite tax credits → the flywheel accelerates",
]
for i, step in enumerate(flywheel_steps):
    add_bullet(doc, step, bold_prefix=f"Step {i+1}: ")

doc.add_heading('5.5 Projected Impact (Year 1 — 1,000 businesses)', level=2)

add_table(doc,
    ["Metric", "Value"],
    [
        ["Tax credits granted", "2,000,000,000 CFA ($3.1M)"],
        ["Data points generated", "120,000 reports → 500,000 enriched points"],
        ["Revenue from data sales", "3,250,000,000 CFA ($5.0M)"],
        ["NET GAIN FOR THE STATE", "+1,250,000,000 CFA (+$1.9M)"],
    ],
    col_widths=[7, 8]
)

add_paragraph(doc, "The State earns 62% more than it gave away in tax credits.")

doc.add_heading('5.6 Anti-Fraud Protections', level=2)

add_paragraph(doc, "Any false data declarations are subject to General Tax Code fraud penalties, increased by 50%. Business data is cross-validated against citizen declarations: if a company claims sales that citizen reporters contradict, the declaration is automatically flagged for review.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#            6. PILLAR 3: DIGITAL FASO MEABO
# ═══════════════════════════════════════════════════════════════

doc.add_heading('6. Pillar 3: Digital Faso Meabo', level=1)
add_horizontal_line(doc)

doc.add_heading('6.1 The Problem with Public Contracts', level=2)

add_paragraph(doc, "Public contracts in Burkina Faso suffer from three chronic issues: delays (projects regularly exceed their deadline by 50–100%), cost overruns (frequent amendments increasing the budget by 20–30%), and no citizen oversight (beneficiary communities have no visibility into project progress).")

doc.add_heading('6.2 Tokenized Milestone Payment', level=2)

add_bold_paragraph(doc, "Old System:", size=12)
add_paragraph(doc, "Pay 30% advance regardless of progress → pay 30% more at month 6 → negotiate final payment with lawyers and amendments → project completed late and over budget.")

add_bold_paragraph(doc, "New System:", size=12)
add_paragraph(doc, "Each project is divided into milestone segments. Each segment = 1 milestone token. Token is released ONLY when: (1) site foreman declares segment complete, (2) citizen monitor confirms, (3) GPS-timestamped photo submitted, (4) automated quality check passes. 3 out of 4 confirmations trigger automatic Treasury payment. No official signature needed. No inspection mission. No corruption opportunity.")

doc.add_heading('6.3 Speed Incentives', level=2)

add_table(doc,
    ["Situation", "Consequence"],
    [
        ["Completed 4 months early", "+10% bonus"],
        ["Completed 2 months early", "+5% bonus"],
        ["Completed on time", "Base payment"],
        ["1 month late", "–2% per month automatic penalty"],
        ["6 months late", "Automatic 2-year blacklist"],
    ],
    col_widths=[7, 8]
)

doc.add_heading('6.4 Citizen Verification', level=2)

add_paragraph(doc, "Each public contract is monitored by citizen monitors — residents living near the worksite who report daily via USSD. Reports include: work progressing normally, work stopped, workers absent, equipment idle, quality concerns. Payment: 200 CFA per report. Citizen reports cross-validate contractor claims.")

doc.add_heading('6.5 Community Labor (Faso Meabo)', level=2)

add_paragraph(doc, "In Burkinabè tradition, the community mobilizes for shared public works. Digital Faso Meabo formalizes and compensates this tradition:")

add_table(doc,
    ["Task", "Difficulty", "Daily Token", "Monthly (22 days)"],
    [
        ["Vegetation clearing", "Light", "1,500 CFA", "33,000 CFA"],
        ["Earthwork / rock removal", "Medium", "2,000 CFA", "44,000 CFA"],
        ["Drainage ditch digging", "Medium", "2,000 CFA", "44,000 CFA"],
        ["Manual material transport", "Heavy", "2,500 CFA", "55,000 CFA"],
        ["Masonry assistance", "Skilled", "3,000 CFA", "66,000 CFA"],
        ["Site coordination", "Leadership", "4,000 CFA", "88,000 CFA"],
    ],
    col_widths=[4.5, 3, 3.5, 4]
)

add_paragraph(doc, "In each contract, the unskilled labor portion (15–20% of budget) is ring-fenced for Faso Meabo. Villagers provide the labor and are paid directly via mobile money. The contractor benefits from free labor. The community earns income and has a stake in project quality.")

doc.add_heading('6.6 Anti-Corruption Impact', level=2)

add_table(doc,
    ["Corruption Type", "Current Situation", "With WASI"],
    [
        ["Ghost workers", "Bill 100, only 30 exist", "USSD check-in with GPS — can't fake 70 phones"],
        ["Fake progress", "\"60% done\" when it's 30%", "12 citizens contradict daily + photos required"],
        ["Inflated invoices", "\"500 tonnes cement used\"", "Cross-check with suppliers' tokenized data"],
        ["Kickbacks", "Official signs for 10% cut", "No signature needed — algorithm pays"],
        ["Abusive amendments", "Budget doubled via changes", "Milestone map is immutable on platform"],
    ],
    col_widths=[3.5, 5, 7.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#             7. NATIONAL SECURITY IMPACT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('7. National Security Impact', level=1)
add_horizontal_line(doc)

doc.add_heading('7.1 Terrorism is a Price Problem', level=2)

add_paragraph(doc, "In the Sahel, Centre-Nord, Est, Nord, and Boucle du Mouhoun regions, armed groups (JNIM, ISGS) recruit primarily among idle youth. The mechanism is economic:")

add_table(doc,
    ["", "Armed Group Offer", "State Offer (Today)", "WASI Offer"],
    [
        ["Monthly income", "15,000 CFA + motorcycle", "0 CFA", "14,000–82,000 CFA"],
        ["Risk", "Death, family targeted", "N/A", "Zero"],
        ["Dignity", "Criminal status", "Abandonment", "Connected to the State"],
        ["Future", "None", "None", "Career progression, micro-credit access"],
    ],
    col_widths=[3.5, 4, 4, 4.5]
)

doc.add_heading('7.2 Cost Comparison', level=2)

add_table(doc,
    ["Item", "Annual Cost"],
    [
        ["WASI UBDI program (200,000 participants)", "14,400,000,000 CFA (~$22M)"],
        ["", ""],
        ["Cost of terrorism to Burkina Faso:", ""],
        ["  Military spending", "325,000,000,000 CFA"],
        ["  Economic losses (displacement, trade)", "1,300,000,000,000 CFA"],
        ["  Humanitarian response", "585,000,000,000 CFA"],
        ["  Lost foreign investment", "650,000,000,000 CFA"],
        ["  Total cost of terrorism", "~2,860,000,000,000 CFA (~$4.4B)"],
    ],
    col_widths=[9, 7]
)

p = add_paragraph(doc, "The program costs 0.5% of what terrorism costs. Even a 10% reduction in recruitment yields a 20:1 ROI.")
for run in p.runs:
    run.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

doc.add_heading('7.3 The Early Warning System', level=2)

add_paragraph(doc, "Daily declarations create a human sensor network covering fragile zones:")

add_bullet(doc, " = zone is stable, no intervention needed", bold_prefix="Normal presence")
add_bullet(doc, " = event to monitor (market day? network outage?)", bold_prefix="20% drop")
add_bullet(doc, " = automatic security alert → CONASUR, OCHA, Defense Forces", bold_prefix="50%+ drop")
add_bullet(doc, " = probable emergency situation", bold_prefix="Total silence")

add_paragraph(doc, "This system costs zero in equipment (phones already exist) and provides granularity that neither satellites nor drones can achieve.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#              8. TECHNICAL ARCHITECTURE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('8. Technical Architecture', level=1)
add_horizontal_line(doc)

doc.add_heading('8.1 Simplified Overview', level=2)

add_table(doc,
    ["Layer", "Components"],
    [
        ["Data Sources", "Citizens (USSD), Businesses (USSD/Web), Scrapers (World Bank, IMF, ACLED, Comtrade, Pink Sheet), Government databases"],
        ["Gateway", "Africa's Talking / Orange Business Services USSD gateway"],
        ["Platform Core", "Collection & validation engine, AI anti-fraud, Aggregation & anonymization, Data sales API, Real-time dashboard, Payment engine"],
        ["Outputs", "Mobile money payments to citizens, Secure API for global data buyers (AI agents, fintech, traders, NGOs, governments)"],
    ],
    col_widths=[4, 12]
)

doc.add_heading('8.2 Data Security', level=2)

add_table(doc,
    ["Protection", "Implementation"],
    [
        ["Anonymization", "All phone numbers SHA-256 hashed"],
        ["Aggregation", "No individual data sold — minimum 10 persons per aggregate"],
        ["Compliance", "CIL (Burkina Faso Data Protection Authority) authorization"],
        ["Sovereignty", "Servers hosted in Burkina Faso or ECOWAS zone"],
        ["Encryption", "All communications via TLS 1.3"],
    ],
    col_widths=[4, 12]
)

doc.add_heading('8.3 Existing Infrastructure', level=2)

add_paragraph(doc, "The WASI platform is already operational with: a WASI index calculation engine covering all 16 ECOWAS countries, a functional USSD module, 5 real-time data scrapers, a credit-based billing system, AI-powered data quality guardrails (4 levels), and a banking module with credit scoring. The technical delta for the full program is estimated at 4–6 months of development.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#     9. ECONOMIC MODEL AND FINANCIAL PROJECTIONS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('9. Economic Model and Financial Projections', level=1)
add_horizontal_line(doc)

doc.add_heading('9.1 Revenue Sources for the State', level=2)

add_bold_paragraph(doc, "Direct Revenue (CFA)", size=12)

add_table(doc,
    ["Source", "Year 1", "Year 2", "Year 3"],
    [
        ["Sovereign share of data sales (30%)", "350M", "1,400M", "4,680M"],
        ["Additional tax receipts (DTTC formalization)", "480M", "2,800M", "8,000M"],
        ["Savings on public contracts (Faso Meabo)", "200M", "1,000M", "3,000M"],
        ["Total direct revenue", "1,030M", "5,200M", "15,680M"],
    ],
    col_widths=[6, 3, 3, 3]
)

add_bold_paragraph(doc, "Costs to the State (CFA)", size=12)

add_table(doc,
    ["Cost", "Year 1", "Year 2", "Year 3"],
    [
        ["DTTC tax credits granted", "500M", "2,000M", "3,750M"],
        ["Citizen payment subsidy", "200M", "600M", "1,200M"],
        ["Administration and oversight", "100M", "200M", "300M"],
        ["Total costs", "800M", "2,800M", "5,250M"],
    ],
    col_widths=[6, 3, 3, 3]
)

add_bold_paragraph(doc, "Net Benefit to the State", size=14, color=RGBColor(0x00, 0x6B, 0x3F))

add_table(doc,
    ["", "Year 1", "Year 2", "Year 3"],
    [
        ["Net benefit (CFA)", "+230,000,000", "+2,400,000,000", "+10,430,000,000"],
        ["Net benefit (USD)", "+$354,000", "+$3.7M", "+$16.0M"],
    ],
    col_widths=[5, 3.5, 3.5, 3.5]
)

doc.add_heading('9.2 Financial Flow Structure', level=2)

add_paragraph(doc, "For every 100 CFA of data sold to international buyers:")

add_table(doc,
    ["Recipient", "Share", "Amount"],
    [
        ["State of Burkina Faso (sovereign royalty)", "30%", "30 CFA"],
        ["Citizen or business (producer compensation)", "40%", "40 CFA"],
        ["WASI (platform operations)", "20%", "20 CFA"],
        ["Village Data Chief (local quality assurance)", "10%", "10 CFA"],
    ],
    col_widths=[7, 3, 4]
)

doc.add_heading('9.3 Initial Investment Required', level=2)

add_table(doc,
    ["Item", "CFA", "USD"],
    [
        ["Technical development", "650,000,000", "$1,000,000"],
        ["Server and USSD infrastructure", "400,000,000", "$615,000"],
        ["Data Chief recruitment and training", "500,000,000", "$770,000"],
        ["Awareness campaign and registration", "300,000,000", "$460,000"],
        ["Working capital (mobile payments)", "850,000,000", "$1,310,000"],
        ["Administration and legal", "200,000,000", "$310,000"],
        ["Contingency (10%)", "300,000,000", "$460,000"],
        ["TOTAL", "3,200,000,000", "$4,925,000"],
    ],
    col_widths=[7, 4.5, 3.5]
)

add_paragraph(doc, "Potential funding sources: national budget, Sahel Alliance (EU), UNDP Stabilization Facility, AFD, GIZ, World Bank (IDA), Bill & Melinda Gates Foundation.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#              10. PROPOSED LEGAL FRAMEWORK
# ═══════════════════════════════════════════════════════════════

doc.add_heading('10. Proposed Legal Framework', level=1)
add_horizontal_line(doc)

doc.add_heading('10.1 Decree Establishing the Program', level=2)

articles = [
    ("Article 1 — Purpose.", " A National Sovereign Data Tokenization Program (NSDTP) is hereby established for the collection, verification, aggregation, and commercialization of economic data produced by citizens and businesses of Burkina Faso."),
    ("Article 2 — Oversight.", " The NSDTP is placed under the joint oversight of the Ministry of Digital Transition and the Ministry of Economy and Finance."),
    ("Article 3 — Technical operator.", " Technical operation of the NSDTP is entrusted by concession contract to a private operator selected according to the Public Procurement Code, for an initial period of five (5) years, renewable."),
    ("Article 4 — Data sovereignty.", " Data collected under the NSDTP remains the sovereign property of the State of Burkina Faso. The technical operator has the right to operate and commercialize, under State oversight."),
    ("Article 5 — Privacy protection.", " All personal data is anonymized in accordance with the personal data protection law. No individual data may be sold or transmitted to a third party."),
]

for bold, rest in articles:
    add_paragraph(doc, rest, bold_prefix=bold)

doc.add_heading('10.2 Finance Law Amendment (DTTC)', level=2)

dttc_articles = [
    ("Article X — Definition.", " The DTTC is granted to registered businesses that contribute structured economic data to the NSDTP."),
    ("Article X+1 — Cap.", " The DTTC may not exceed 25% of the tax owed, with an absolute cap of 5,000,000 CFA per business per year."),
    ("Article X+2 — Conditions.", " Data must be: (a) truthful and verifiable; (b) transmitted on schedule; (c) compliant with platform standards; (d) validated by the automated quality control system."),
    ("Article X+3 — Penalties.", " False data declarations are subject to tax fraud penalties, increased by 50%."),
]

for bold, rest in dttc_articles:
    add_paragraph(doc, rest, bold_prefix=bold)

doc.add_heading('10.3 Regulatory Authorizations', level=2)

add_table(doc,
    ["Authority", "Authorization", "Purpose"],
    [
        ["CIL (Data Protection Commission)", "Processing authorization", "Data collection and commercialization"],
        ["ARCEP", "USSD license", "Use of telecom infrastructure"],
        ["DGMP (Public Procurement)", "Convention", "Public contract monitoring"],
        ["BCEAO", "Notification", "Mobile money payment flows"],
    ],
    col_widths=[5, 4.5, 5.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#             11. IMPLEMENTATION TIMELINE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('11. Implementation Timeline', level=1)
add_horizontal_line(doc)

doc.add_heading('Phase 1: Pilot (Months 1–6)', level=2)

add_table(doc,
    ["Month", "Activity"],
    [
        ["1", "Sign MOU with ministries"],
        ["1–2", "Obtain CIL and ARCEP authorizations"],
        ["2–3", "Complete technical development"],
        ["3–4", "Recruit and train 20 Data Chiefs in 4 provinces (Kadiogo, Houet, Mouhoun, Nahouri)"],
        ["4–5", "Register 500 citizen reporters + 50 businesses"],
        ["5–6", "Launch pilot — first payments, first data buyers"],
    ],
    col_widths=[3, 13]
)

add_bold_paragraph(doc, "Target: 15,000 data points/month, 5 pilot buyers", size=11)

doc.add_heading('Phase 2: National Rollout (Months 7–18)', level=2)

add_table(doc,
    ["Month", "Activity"],
    [
        ["7–9", "Expand to all 13 regions"],
        ["7–12", "Register 50,000 citizens + 2,000 businesses"],
        ["9–12", "Launch DTTC (next fiscal year)"],
        ["10–14", "Integrate Faso Meabo on 10 pilot public contracts"],
        ["12–18", "Scale to 200,000 citizens, 10,000 businesses"],
    ],
    col_widths=[3, 13]
)

add_bold_paragraph(doc, "Target: 5,000,000 data points/month, 60 buyers", size=11)

doc.add_heading('Phase 3: ECOWAS Expansion (Months 19–36)', level=2)

add_table(doc,
    ["Month", "Activity"],
    [
        ["19–24", "Replicate in Côte d'Ivoire and Senegal"],
        ["24–30", "Expand to Ghana and Mali"],
        ["30–36", "Coverage of 6+ ECOWAS countries"],
    ],
    col_widths=[3, 13]
)

add_bold_paragraph(doc, "Target: Burkina Faso at the center of Africa's largest real-time economic data network.", size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#            12. INTERNATIONAL PARTNERSHIPS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('12. International Partnerships', level=1)
add_horizontal_line(doc)

doc.add_heading('12.1 Funding Partners', level=2)

add_table(doc,
    ["Partner", "Target Amount", "Instrument"],
    [
        ["Sahel Alliance (EU)", "1,500,000,000 CFA", "Stabilization grant"],
        ["UNDP Stabilization Facility", "650,000,000 CFA", "Grant"],
        ["AFD (France)", "500,000,000 CFA", "Concessional loan"],
        ["GIZ (Germany)", "300,000,000 CFA", "Technical cooperation"],
        ["World Bank (IDA)", "800,000,000 CFA", "Credit"],
        ["Gates Foundation", "400,000,000 CFA", "Grant"],
    ],
    col_widths=[5, 5.5, 4.5]
)

doc.add_heading('12.2 Strategic Buyers', level=2)

add_table(doc,
    ["Buyer", "Interest", "Target Contract"],
    [
        ["WFP (World Food Programme)", "Real-time food prices", "$300,000/year"],
        ["UNDSS (UN Security)", "Stability mapping", "$500,000/year"],
        ["World Bank", "Economic indicators", "$200,000/year"],
        ["Anthropic / OpenAI", "AI training data", "$100,000–$500,000/year"],
        ["ECOWAS", "Cross-border trade flows", "Revenue sharing"],
    ],
    col_widths=[5, 5.5, 4.5]
)

doc.add_heading('12.3 Diplomatic Dimension', level=2)

add_paragraph(doc, "By launching this program, Burkina Faso positions itself as:")

diplomatic_points = [
    "The first African country to monetize its sovereign data for the benefit of its population",
    "A leader in digital sovereignty on the continent",
    "A model for ECOWAS and the African Union",
    "A partner of choice for AI companies seeking ethically sourced African data",
]
for pt in diplomatic_points:
    add_bullet(doc, pt)

p = add_paragraph(doc, "His Excellency the President of Faso could announce this initiative at the next African Union Summit as a major Burkinabè contribution to Agenda 2063.")
for run in p.runs:
    run.bold = True
    run.font.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#                      13. ANNEXES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('13. Annexes', level=1)
add_horizontal_line(doc)

doc.add_heading('Annex A: WASI Index Country Weights', level=2)

add_table(doc,
    ["Category", "Country", "Code", "Weight"],
    [
        ["Primary (75%)", "Nigeria", "NG", "28%"],
        ["", "Côte d'Ivoire", "CI", "22%"],
        ["", "Ghana", "GH", "15%"],
        ["", "Senegal", "SN", "10%"],
        ["Secondary (20%)", "Burkina Faso", "BF", "4%"],
        ["", "Mali", "ML", "4%"],
        ["", "Guinea", "GN", "4%"],
        ["", "Benin", "BJ", "3%"],
        ["", "Togo", "TG", "3%"],
        ["Tertiary (5%)", "Niger", "NE", "1%"],
        ["", "Mauritania", "MR", "1%"],
        ["", "Guinea-Bissau", "GW", "1%"],
        ["", "Sierra Leone", "SL", "1%"],
        ["", "Liberia", "LR", "1%"],
        ["", "Gambia", "GM", "1%"],
        ["", "Cape Verde", "CV", "1%"],
    ],
    col_widths=[4, 4, 2.5, 2.5]
)

doc.add_heading('Annex B: Data Protection Guarantees', level=2)

add_table(doc,
    ["Principle", "Implementation"],
    [
        ["Anonymization", "Phone numbers hashed SHA-256"],
        ["Aggregation", "Minimum 10 persons per aggregate"],
        ["Consent", "Explicit acceptance during USSD registration"],
        ["Right of access", "Citizen can view data via *384*WASI*7#"],
        ["Right to deletion", "Citizen can request account deletion"],
        ["Data localization", "Servers in Burkina Faso or ECOWAS zone"],
        ["Audit", "Annual report to CIL"],
    ],
    col_widths=[4, 12]
)

# ── Footer ──────────────────────────────────────────────────
doc.add_paragraph()
add_horizontal_line(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("This document is the property of WASI. Reproduction without authorization is prohibited.")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Reference: WASI/BF/GOV/2026-001 — March 2026")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'

# ── Save ────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "WASI_Government_Proposal_BF_EN.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
